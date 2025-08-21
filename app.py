# app.py - FastAPI Backend für Company Research Tool

import asyncio
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
import logging

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel

# Import existing LangGraph components
from src.agent.graph import graph
from src.agent.state import EXTENDED_EXTRACTION_SCHEMA
from src.agent.debug_utils import analyze_data_quality

# PDF/Word export libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Company Research Tool", version="1.0.0")

# No static files needed - everything is embedded

# In-memory storage für aktive Jobs (in production: Redis/DB)
active_jobs: Dict[str, Dict[str, Any]] = {}

# Pydantic Models
class ResearchRequest(BaseModel):
    urls: list[str]
    generate_cold_email: bool = True
    email_config: Dict[str, str] = {
        "sender_company": "Mobile Fusion",
        "sender_name": "Jonas Kremser",
        "sender_role": "Digital Marketing Consultant",
        "service_offering": "SEO & Meta Ad Optimierung",
        "email_tone": "professionell",
        "email_length": "medium",
        "call_to_action": "kostenloses Beratungsgespräch",
        "email_language": "deutsch"
    }
    user_notes: str = "Fokus auf SEO und Meta Advertising"

class JobStatus(BaseModel):
    job_id: str
    status: str  # "started", "crawling", "analyzing", "generating_emails", "completed", "error"
    progress: int  # 0-100
    current_step: str
    results: Optional[Dict[str, Any]] = None
    error_message: Optional[str] = None
    error_category: Optional[str] = None
    error_details: Optional[str] = None

# WebSocket Manager für Live-Updates
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}

    async def connect(self, websocket: WebSocket, job_id: str):
        await websocket.accept()
        self.active_connections[job_id] = websocket

    def disconnect(self, job_id: str):
        if job_id in self.active_connections:
            del self.active_connections[job_id]

    async def send_update(self, job_id: str, message: dict):
        if job_id in self.active_connections:
            try:
                await self.active_connections[job_id].send_json(message)
            except Exception as e:
                logger.error(f"Error sending WebSocket message: {e}")
                self.disconnect(job_id)

manager = ConnectionManager()

# Helper Functions
def update_job_progress(job_id: str, status: str, progress: int, current_step: str, results: Optional[Dict] = None):
    """Update job progress and send WebSocket notification"""
    if job_id in active_jobs:
        active_jobs[job_id].update({
            "status": status,
            "progress": progress,
            "current_step": current_step,
            "results": results,
            "updated_at": datetime.now().isoformat()
        })
        
        # Send WebSocket update
        asyncio.create_task(manager.send_update(job_id, {
            "job_id": job_id,
            "status": status,
            "progress": progress,
            "current_step": current_step,
            "results": results
        }))

async def run_research_analysis(job_id: str, request: ResearchRequest):
    """Run the complete research analysis with progress updates"""
    try:
        update_job_progress(job_id, "started", 10, "Initialisierung...")
        
        # Prepare state for LangGraph
        state = {
            "urls": request.urls,
            "extraction_schema": EXTENDED_EXTRACTION_SCHEMA,
            "user_notes": request.user_notes,
            "completed_notes": [],
            "info": {},
            "is_satisfactory": {},
            "reflection_steps_taken": {},
            "meta_ad_intelligence": {},
            "generate_cold_email": request.generate_cold_email,
            "email_config": request.email_config,
            "generated_emails": {}
        }
        
        update_job_progress(job_id, "crawling", 25, "Website-Crawling läuft...")
        
        # Run LangGraph analysis
        result = await graph.ainvoke(state)
        
        update_job_progress(job_id, "analyzing", 60, "Meta Ad Intelligence Analysis...")
        
        # Add quality analysis
        quality_reports = {}
        for url in request.urls:
            if url in result.get('info', {}):
                website_data = result['info'][url]
                meta_data = result.get('meta_ad_intelligence', {}).get(url)
                quality_reports[url] = analyze_data_quality(website_data, meta_data)
        
        result['quality_reports'] = quality_reports
        
        update_job_progress(job_id, "generating_emails", 85, "E-Mails werden generiert...")
        
        update_job_progress(job_id, "completed", 100, "Analyse abgeschlossen!", result)
        
    except Exception as e:
        logger.error(f"Error in research analysis: {e}")
        
        # Categorize error types for better user feedback
        error_message = str(e)
        error_category = "Allgemeiner Fehler"
        error_details = f"Technische Details: {error_message}"
        
        if "META_API" in error_message or "Token" in error_message:
            error_category = "Meta API Konfigurationsfehler"
            error_message = "Meta API Token ist ungültig oder abgelaufen. Bitte überprüfen Sie Ihre .env Datei."
            error_details = f"Ursprünglicher Fehler: {str(e)}"
        elif "FIRECRAWL" in error_message or "crawl" in error_message.lower():
            error_category = "Website-Crawling Fehler"
            error_message = "Firecrawl API konnte die Website nicht laden. Überprüfen Sie den API-Key oder die URL."
            error_details = f"Ursprünglicher Fehler: {str(e)}"
        elif "OPENAI" in error_message or "openai" in error_message.lower():
            error_category = "OpenAI API Fehler"
            error_message = "OpenAI API Problem. Überprüfen Sie Ihren API-Key und das Guthaben."
            error_details = f"Ursprünglicher Fehler: {str(e)}"
        elif "timeout" in error_message.lower():
            error_category = "Timeout Fehler"
            error_message = "Die Analyse hat zu lange gedauert. Versuchen Sie es mit einer einfacheren Website."
        
        active_jobs[job_id]["status"] = "error"
        active_jobs[job_id]["error_message"] = error_message
        active_jobs[job_id]["error_category"] = error_category
        active_jobs[job_id]["error_details"] = error_details
        
        await manager.send_update(job_id, {
            "job_id": job_id,
            "status": "error",
            "error_message": error_message,
            "error_category": error_category,
            "error_details": error_details
        })

# API Endpoints
@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the main HTML page"""
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Company Research Tool</title>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1">
        <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
        <style>
            .progress-step { margin: 10px 0; }
            .result-section { margin-top: 20px; }
            .company-card { border: 1px solid #ddd; padding: 15px; margin: 10px 0; border-radius: 5px; }
            .export-buttons { margin-top: 15px; }
        </style>
    </head>
    <body>
        <div class="container mt-5">
            <h1 class="mb-4">🔍 Company Research Tool</h1>
            
            <!-- Input Form -->
            <div class="card mb-4">
                <div class="card-header">
                    <h5>Neue Analyse starten</h5>
                </div>
                <div class="card-body">
                    <form id="researchForm">
                        <div class="mb-3">
                            <label for="urls" class="form-label">Website URLs (eine pro Zeile)</label>
                            <textarea class="form-control" id="urls" rows="3" placeholder="https://example.com&#10;https://another-company.com" required></textarea>
                        </div>
                        
                        <div class="row">
                            <div class="col-md-6">
                                <label for="senderName" class="form-label">Ihr Name</label>
                                <input type="text" class="form-control" id="senderName" value="Jonas Kremser">
                            </div>
                            <div class="col-md-6">
                                <label for="senderCompany" class="form-label">Ihr Unternehmen</label>
                                <input type="text" class="form-control" id="senderCompany" value="Mobile Fusion">
                            </div>
                        </div>
                        
                        <div class="row mt-3">
                            <div class="col-md-6">
                                <label for="senderRole" class="form-label">Ihre Position</label>
                                <input type="text" class="form-control" id="senderRole" value="Digital Marketing Consultant">
                            </div>
                            <div class="col-md-6">
                                <label for="serviceOffering" class="form-label">Service-Angebot</label>
                                <input type="text" class="form-control" id="serviceOffering" value="SEO & Meta Ad Optimierung">
                            </div>
                        </div>
                        
                        <div class="mt-3">
                            <label for="userNotes" class="form-label">Zusätzliche Notizen</label>
                            <textarea class="form-control" id="userNotes" rows="2" placeholder="Fokus auf digitales Marketing..."></textarea>
                        </div>
                        
                        <button type="submit" class="btn btn-primary mt-3">🚀 Analyse starten</button>
                    </form>
                </div>
            </div>
            
            <!-- Progress Section -->
            <div id="progressSection" style="display: none;" class="card mb-4">
                <div class="card-header">
                    <h5>Fortschritt</h5>
                </div>
                <div class="card-body">
                    <div class="progress mb-3">
                        <div id="progressBar" class="progress-bar" role="progressbar" style="width: 0%">0%</div>
                    </div>
                    <div id="currentStep">Warten...</div>
                    
                    <!-- Error Display -->
                    <div id="errorAlert" class="alert alert-danger mt-3" style="display: none;">
                        <h6 class="alert-heading">❌ Fehler aufgetreten:</h6>
                        <p id="errorMessage"></p>
                        <hr>
                        <div id="errorDetails" style="display: none;">
                            <small id="errorDetailsText"></small>
                        </div>
                        <button class="btn btn-outline-danger btn-sm mt-2" onclick="showErrorDetails()">Details anzeigen</button>
                        <button class="btn btn-primary btn-sm mt-2 ms-2" onclick="resetForm()">Neue Analyse starten</button>
                    </div>
                </div>
            </div>
            
            <!-- Results Section -->
            <div id="resultsSection" style="display: none;" class="card">
                <div class="card-header d-flex justify-content-between align-items-center">
                    <h5>Ergebnisse</h5>
                    <div class="export-buttons">
                        <button id="exportPDF" class="btn btn-outline-danger btn-sm">📄 PDF Export</button>
                        <button id="exportWord" class="btn btn-outline-primary btn-sm">📝 Word Export</button>
                    </div>
                </div>
                <div class="card-body">
                    <div id="resultsContent"></div>
                </div>
            </div>
        </div>
        
        <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
        <script>
            let currentJobId = null;
            let currentResults = null;
            
            // Form submission
            document.getElementById('researchForm').addEventListener('submit', async (e) => {
                e.preventDefault();
                
                const urls = document.getElementById('urls').value.split('\\n').filter(url => url.trim());
                const emailConfig = {
                    sender_name: document.getElementById('senderName').value,
                    sender_company: document.getElementById('senderCompany').value,
                    sender_role: document.getElementById('senderRole').value,
                    service_offering: document.getElementById('serviceOffering').value,
                    email_tone: "professionell",
                    email_length: "medium",
                    call_to_action: "kostenloses Beratungsgespräch",
                    email_language: "deutsch"
                };
                
                const requestData = {
                    urls: urls,
                    generate_cold_email: true,
                    email_config: emailConfig,
                    user_notes: document.getElementById('userNotes').value || "Fokus auf SEO und Meta Advertising"
                };
                
                try {
                    const response = await fetch('/api/research', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(requestData)
                    });
                    
                    if (!response.ok) {
                        const errorData = await response.json();
                        throw new Error(`Server Fehler (${response.status}): ${errorData.detail || 'Unbekannter Fehler'}`);
                    }
                    
                    const result = await response.json();
                    currentJobId = result.job_id;
                    
                    // Show progress section and clear any previous errors
                    document.getElementById('progressSection').style.display = 'block';
                    document.getElementById('resultsSection').style.display = 'none';
                    document.getElementById('errorAlert').style.display = 'none';
                    
                    // Reset progress bar
                    document.getElementById('progressBar').className = 'progress-bar';
                    
                    // Connect WebSocket
                    connectWebSocket(currentJobId);
                    
                } catch (error) {
                    console.error('Request error:', error);
                    
                    // Make sure progress section is visible for error display
                    document.getElementById('progressSection').style.display = 'block';
                    showError('Verbindungsfehler', error.message);
                }
            });
            
            // WebSocket connection
            function connectWebSocket(jobId) {
                const ws = new WebSocket(`ws://localhost:8000/ws/${jobId}`);
                
                ws.onmessage = (event) => {
                    const data = JSON.parse(event.data);
                    updateProgress(data);
                };
                
                ws.onerror = (error) => {
                    console.error('WebSocket error:', error);
                    showError('Verbindungsfehler', 'WebSocket-Verbindung unterbrochen. Eventuell ist der Server nicht erreichbar.');
                };
                
                ws.onclose = (event) => {
                    if (event.code !== 1000) { // 1000 = normal closure
                        console.warn('WebSocket closed unexpectedly:', event);
                        showError('Verbindung verloren', 'Die Verbindung zum Server wurde unterbrochen.');
                    }
                };
            }
            
            // Update progress
            function updateProgress(data) {
                const progressBar = document.getElementById('progressBar');
                const currentStep = document.getElementById('currentStep');
                
                progressBar.style.width = data.progress + '%';
                progressBar.textContent = data.progress + '%';
                currentStep.textContent = data.current_step;
                
                if (data.status === 'completed') {
                    progressBar.className = 'progress-bar bg-success';
                    currentResults = data.results;
                    showResults(data.results);
                } else if (data.status === 'error') {
                    progressBar.className = 'progress-bar bg-danger';
                    showError(
                        data.error_category || 'Analyse-Fehler', 
                        data.error_message || 'Unbekannter Fehler',
                        data.error_details || null
                    );
                }
            }
            
            // Show error function
            function showError(title, message, details = null) {
                const errorAlert = document.getElementById('errorAlert');
                const errorMessage = document.getElementById('errorMessage');
                const errorDetails = document.getElementById('errorDetailsText');
                const currentStep = document.getElementById('currentStep');
                
                // Set error message
                errorMessage.innerHTML = `<strong>${title}:</strong> ${message}`;
                
                // Set details if available
                if (details) {
                    errorDetails.textContent = details;
                    document.getElementById('errorDetails').style.display = 'block';
                }
                
                // Show error alert
                errorAlert.style.display = 'block';
                currentStep.textContent = `❌ ${title}`;
                
                // Scroll to error for better visibility
                errorAlert.scrollIntoView({ behavior: 'smooth', block: 'nearest' });
                
                // Categorize error types
                if (message.includes('Meta') || message.includes('TOKEN') || message.includes('API')) {
                    errorMessage.innerHTML += `<br><small class="text-muted">💡 Tipp: Überprüfen Sie Ihre API-Konfiguration (.env Datei)</small>`;
                } else if (message.includes('Verbindung') || message.includes('WebSocket')) {
                    errorMessage.innerHTML += `<br><small class="text-muted">💡 Tipp: Server neu starten mit 'python start_tool.py'</small>`;
                } else if (message.includes('crawl') || message.includes('Firecrawl')) {
                    errorMessage.innerHTML += `<br><small class="text-muted">💡 Tipp: Firecrawl API-Key überprüfen</small>`;
                }
            }
            
            // Show error details
            function showErrorDetails() {
                const details = document.getElementById('errorDetails');
                const button = event.target;
                
                if (details.style.display === 'none') {
                    details.style.display = 'block';
                    button.textContent = 'Details ausblenden';
                } else {
                    details.style.display = 'none';
                    button.textContent = 'Details anzeigen';
                }
            }
            
            // Reset form
            function resetForm() {
                document.getElementById('progressSection').style.display = 'none';
                document.getElementById('resultsSection').style.display = 'none';
                document.getElementById('errorAlert').style.display = 'none';
                document.getElementById('progressBar').className = 'progress-bar';
                document.getElementById('progressBar').style.width = '0%';
                document.getElementById('progressBar').textContent = '0%';
                currentJobId = null;
                currentResults = null;
            }
            
            // Show results
            function showResults(results) {
                document.getElementById('resultsSection').style.display = 'block';
                const content = document.getElementById('resultsContent');
                
                let html = '';
                
                // Website Analysis Results
                if (results.info) {
                    for (const [url, data] of Object.entries(results.info)) {
                        html += `
                            <div class="company-card">
                                <h6>🏢 ${data.company_name || 'Unbekannt'}</h6>
                                <p><strong>URL:</strong> ${url}</p>
                                <p><strong>USP:</strong> ${data.unique_selling_proposition || 'N/A'}</p>
                                <p><strong>UX Rating:</strong> ${data.website_user_experience?.overall_ux_rating || 'N/A'}</p>
                            </div>
                        `;
                    }
                }
                
                // Meta Ad Intelligence
                if (results.meta_ad_intelligence) {
                    html += '<h6>📱 Meta Ad Intelligence:</h6>';
                    for (const [url, data] of Object.entries(results.meta_ad_intelligence)) {
                        const llm = data.llm_analysis || {};
                        html += `
                            <div class="company-card">
                                <p><strong>Status:</strong> ${llm.advertising_status || 'N/A'}</p>
                                <p><strong>Budget:</strong> ${llm.budget_assessment || 'N/A'}</p>
                                <p><strong>Opportunities:</strong> ${(llm.optimization_opportunities || []).slice(0,2).join(', ')}</p>
                            </div>
                        `;
                    }
                }
                
                // Generated Emails
                if (results.generated_emails) {
                    html += '<h6>📧 Generierte E-Mails:</h6>';
                    for (const [url, email] of Object.entries(results.generated_emails)) {
                        html += `
                            <div class="company-card">
                                <pre style="white-space: pre-wrap; font-family: inherit;">${email}</pre>
                            </div>
                        `;
                    }
                }
                
                content.innerHTML = html;
            }
            
            // Export functions
            document.getElementById('exportPDF').addEventListener('click', () => {
                if (currentJobId) {
                    window.open(`/api/export/pdf/${currentJobId}`, '_blank');
                }
            });
            
            document.getElementById('exportWord').addEventListener('click', () => {
                if (currentJobId) {
                    window.open(`/api/export/word/${currentJobId}`, '_blank');
                }
            });
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

@app.post("/api/research")
async def start_research(request: ResearchRequest):
    """Start a new research analysis"""
    job_id = str(uuid.uuid4())
    
    # Store job info
    active_jobs[job_id] = {
        "job_id": job_id,
        "status": "queued",
        "progress": 0,
        "current_step": "Warten auf Start...",
        "created_at": datetime.now().isoformat(),
        "request": request.dict(),
        "results": None
    }
    
    # Start analysis in background
    asyncio.create_task(run_research_analysis(job_id, request))
    
    return {"job_id": job_id, "status": "queued"}

@app.get("/api/jobs/{job_id}")
async def get_job_status(job_id: str):
    """Get current job status"""
    if job_id not in active_jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    return active_jobs[job_id]

@app.websocket("/ws/{job_id}")
async def websocket_endpoint(websocket: WebSocket, job_id: str):
    """WebSocket endpoint for live progress updates"""
    await manager.connect(websocket, job_id)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        manager.disconnect(job_id)

@app.get("/api/export/pdf/{job_id}")
async def export_pdf(job_id: str):
    """Export results as PDF"""
    if job_id not in active_jobs or not active_jobs[job_id].get("results"):
        raise HTTPException(status_code=404, detail="Results not found")
    
    results = active_jobs[job_id]["results"]
    
    # Create PDF
    filename = f"company_analysis_{job_id}.pdf"
    filepath = Path(f"exports/{filename}")
    filepath.parent.mkdir(exist_ok=True)
    
    doc = SimpleDocTemplate(str(filepath), pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, spaceAfter=30)
    story.append(Paragraph("Company Research Report", title_style))
    story.append(Spacer(1, 12))
    
    # Add content
    if results.get("info"):
        for url, data in results["info"].items():
            story.append(Paragraph(f"<b>Company: {data.get('company_name', 'Unknown')}</b>", styles['Heading2']))
            story.append(Paragraph(f"URL: {url}", styles['Normal']))
            story.append(Paragraph(f"USP: {data.get('unique_selling_proposition', 'N/A')}", styles['Normal']))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    
    return FileResponse(filepath, filename=filename, media_type="application/pdf")

@app.get("/api/export/word/{job_id}")
async def export_word(job_id: str):
    """Export results as Word document"""
    if job_id not in active_jobs or not active_jobs[job_id].get("results"):
        raise HTTPException(status_code=404, detail="Results not found")
    
    results = active_jobs[job_id]["results"]
    
    # Create Word document
    doc = Document()
    doc.add_heading('Company Research Report', 0)
    
    if results.get("info"):
        for url, data in results["info"].items():
            doc.add_heading(f"Company: {data.get('company_name', 'Unknown')}", level=1)
            doc.add_paragraph(f"URL: {url}")
            doc.add_paragraph(f"USP: {data.get('unique_selling_proposition', 'N/A')}")
            
            # Add email if available
            if results.get("generated_emails", {}).get(url):
                doc.add_heading("Generated Email:", level=2)
                doc.add_paragraph(results["generated_emails"][url])
    
    # Save document
    filename = f"company_analysis_{job_id}.docx"
    filepath = Path(f"exports/{filename}")
    filepath.parent.mkdir(exist_ok=True)
    doc.save(str(filepath))
    
    return FileResponse(filepath, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    import uvicorn
    
    # Create exports directory
    Path("exports").mkdir(exist_ok=True)
    
    print("🚀 Starting Company Research Tool...")
    print("💻 Open http://localhost:8000 in your browser")
    
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)